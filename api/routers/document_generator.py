"""
文档生成器模块
负责生成 Word 文档（转写文档、会议纪要）
"""

import os
import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import FILE_CONFIG

logger = logging.getLogger(__name__)


def save_transcript_to_word(transcript_data, filename_prefix="transcript", language="zh", audio_filename=None, file_id=None):
    """将转录结果保存为Word文档"""
    try:
        doc = Document()
        
        # 定义黑色（RGB(0,0,0)）
        black_color = RGBColor(0, 0, 0)
        
        title = doc.add_heading('语音转文字结果', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题为微软雅黑，黑色
        for run in title.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = black_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph()
        
        info_table = doc.add_table(rows=3, cols=2)
        # 恢复原来的表格样式
        info_table.style = 'Light Grid Accent 1'
        
        for row in info_table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)
        
        # 设置表格第一列（标签）为宋体11号加粗，黑色，居中
        info_table.rows[0].cells[0].text = '生成时间'
        label_para = info_table.rows[0].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = False
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置表格第二列（值）为宋体11号加粗，黑色，居中
        info_table.rows[0].cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        value_para = info_table.rows[0].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = False
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[1].cells[0].text = '音频文件'
        label_para = info_table.rows[1].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = False
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[1].cells[1].text = audio_filename or "未知文件"
        value_para = info_table.rows[1].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = False
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        info_table.rows[2].cells[0].text = '文本长度'
        label_para = info_table.rows[2].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = False
        label_run.font.size = Pt(11)
        label_run.font.name = 'SimSun'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        total_chars = sum(len(entry['text']) for entry in transcript_data)
        info_table.rows[2].cells[1].text = f"{total_chars} 字符"
        value_para = info_table.rows[2].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = False
        value_run.font.size = Pt(11)
        value_run.font.name = 'SimSun'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        for entry in transcript_data:
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(entry['speaker'])
            speaker_run.bold = False
            speaker_run.font.size = Pt(12)
            speaker_run.font.name = 'SimSun'
            speaker_run.font.color.rgb = black_color
            speaker_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 设置发言人段落的下间距为0，使内容紧跟在后面
            speaker_para.paragraph_format.space_after = Pt(0)
            
            # 减小发言人和内容的间距，设置段落间距为0
            text_para = doc.add_paragraph()
            text_para.paragraph_format.space_before = Pt(0)
            text_para.paragraph_format.space_after = Pt(0)
            text_run = text_para.add_run(entry['text'])
            text_run.font.size = Pt(12)
            text_run.font.name = 'SimSun'
            text_run.font.color.rgb = black_color
            text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 不同发言人之间的间距保持正常
            doc.add_paragraph()
        
        # ✅ 修复：使用微秒级时间戳 + file_id 确保文件名唯一性
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S_%f')  # 包含微秒
        
        # 如果提供了 file_id，使用前8个字符作为唯一标识
        if file_id:
            file_id_short = file_id.replace('-', '')[:8]  # 移除连字符，取前8位
            filename = f"{filename_prefix}_{timestamp}_{file_id_short}.docx"
        else:
            filename = f"{filename_prefix}_{timestamp}.docx"
        
        filepath = os.path.join(FILE_CONFIG['output_dir'], filename)
        
        doc.save(filepath)
        return filename, filepath
        
    except Exception as e:
        logger.error(f"保存Word文档失败: {e}")
        return None, None


def save_meeting_summary_to_word(transcript_data, summary_data, filename_prefix="meeting_summary", file_id=None, audio_filename=None, audio_duration=None):
    """将会议纪要保存为Word文档"""
    
    def set_table_borders_black(table):
        """设置表格边框为黑色"""
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # 黑色
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
    
    try:
        doc = Document()
        
        # 定义黑色（RGB(0,0,0)）
        black_color = RGBColor(0, 0, 0)
        
        title = doc.add_heading('会议纪要', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置标题为华文中宋 二号字体
        for run in title.runs:
            run.font.name = 'STZhongsong'  # 华文中宋
            run.font.size = Pt(22)  # 二号字体
            run.font.color.rgb = black_color
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '华文中宋')
        
        doc.add_paragraph()
        
        # 计算音频时长（如果未提供，从 transcript_data 计算）
        if audio_duration is None and transcript_data:
            # 从最后一个转写段的结束时间获取总时长
            last_entry = transcript_data[-1] if transcript_data else None
            if last_entry and 'end_time' in last_entry:
                audio_duration = last_entry['end_time']
            else:
                audio_duration = 0
        
        # 格式化时长
        def format_duration(seconds):
            """将秒数格式化为 分钟 秒 的格式"""
            if seconds is None or seconds <= 0:
                return "0分钟 0秒"
            minutes = int(seconds // 60)
            secs = int(seconds % 60)
            return f"{minutes}分钟 {secs}秒"
        
        # 添加生成时间和音频信息
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        # 设置表格边框为黑色
        set_table_borders_black(info_table)
        
        for row in info_table.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(5.0)
        
        # 第一行：生成时间
        info_table.rows[0].cells[0].text = '生成时间'
        label_para = info_table.rows[0].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = False
        label_run.font.size = Pt(16)  # 三号字体
        label_run.font.name = '仿宋_GB2312'  # 仿宋_GB2312
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        label_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        label_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        
        generated_time = summary_data.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        # 转换为更友好的格式：2025年8月31日 下午5:35
        try:
            dt = datetime.strptime(generated_time, '%Y-%m-%d %H:%M:%S')
            formatted_time = dt.strftime('%Y年%m月%d日')
            hour = dt.hour
            minute = dt.minute
            
            # 正确处理上午/下午时间显示
            if hour == 0:
                time_part = "凌晨0"
            elif hour < 12:
                time_part = f"上午{hour}"
            elif hour == 12:
                time_part = "下午12"
            else:
                time_part = f"下午{hour - 12}"
            
            formatted_time = f"{formatted_time} {time_part}:{minute:02d}"
        except Exception as e:
            logger.warning(f"格式化生成时间失败: {e}, 使用原始时间: {generated_time}")
            formatted_time = generated_time
        
        info_table.rows[0].cells[1].text = formatted_time
        value_para = info_table.rows[0].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = False
        value_run.font.size = Pt(16)
        value_run.font.name = '仿宋_GB2312'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        value_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        value_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        
        # 第二行：音频时长
        info_table.rows[1].cells[0].text = '音频时长'
        label_para = info_table.rows[1].cells[0].paragraphs[0]
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_para.runs[0]
        label_run.bold = False
        label_run.font.size = Pt(16)
        label_run.font.name = '仿宋_GB2312'
        label_run.font.color.rgb = black_color
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        label_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        label_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        
        info_table.rows[1].cells[1].text = format_duration(audio_duration)
        value_para = info_table.rows[1].cells[1].paragraphs[0]
        value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = value_para.runs[0]
        value_run.bold = False
        value_run.font.size = Pt(16)
        value_run.font.name = '仿宋_GB2312'
        value_run.font.color.rgb = black_color
        value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
        value_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
        value_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        
        # 第三行：音频文件名（如果有）
        if audio_filename:
            info_table.rows[2].cells[0].text = '音频文件'
            label_para = info_table.rows[2].cells[0].paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = label_para.runs[0]
            label_run.bold = False
            label_run.font.size = Pt(16)
            label_run.font.name = '仿宋_GB2312'
            label_run.font.color.rgb = black_color
            label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            label_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
            label_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
            
            info_table.rows[2].cells[1].text = audio_filename
            value_para = info_table.rows[2].cells[1].paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            value_run = value_para.runs[0]
            value_run.bold = False
            value_run.font.size = Pt(16)
            value_run.font.name = '仿宋_GB2312'
            value_run.font.color.rgb = black_color
            value_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            value_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
            value_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
        else:
            # 如果没有音频文件名，隐藏第三行
            info_table.rows[2].cells[0].text = ''
            info_table.rows[2].cells[1].text = ''
        
        doc.add_paragraph()
        
        # 添加纪要内容，智能识别标题、表格等格式
        raw_text = summary_data.get('raw_text', '')
        lines = raw_text.split('\n')
        
        # 跟踪一级标题和二级标题序号
        current_level1_title = None
        level2_counter = 0
        chinese_numerals = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', 
                           '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                doc.add_paragraph()
                i += 1
                continue
            
            # 如果当前行是"关键词"行
            if line.strip() == '关键词':
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    run.bold = False
                    run.font.name = '仿宋_GB2312'
                    run.font.size = Pt(16)
                    run.font.color.rgb = black_color
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                
                # 检查下一行是否是关键词内容
                if i + 1 < len(lines) and lines[i + 1].strip() and lines[i + 1].strip() != '':
                    i += 1
                    keywords_content = lines[i].strip()
                    if keywords_content:
                        keywords_para = doc.add_paragraph(keywords_content)
                        keywords_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        keywords_para.paragraph_format.space_before = Pt(0)
                        keywords_para.paragraph_format.space_after = Pt(0)
                        keywords_para.paragraph_format.line_spacing = 1.15
                        for run in keywords_para.runs:
                            run.bold = False
                            run.font.name = '仿宋_GB2312'
                            run.font.size = Pt(16)
                            run.font.color.rgb = black_color
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                            run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                
                i += 1
                continue
            
            # 检测表格：查找包含"序号"、"事项描述"等表头的行
            if '序号' in line and ('事项描述' in line or '负责人' in line or '备注' in line):
                # 这是一个表格，开始解析表格
                table_data = []
                headers = []
                
                # 解析表头
                if '|' in line:
                    header_line = line
                    if header_line.startswith('|') and header_line.endswith('|'):
                        headers = [h.strip() for h in header_line.split('|')[1:-1]]
                        i += 1
                        if i < len(lines) and '|' in lines[i] and ('---' in lines[i] or ':-' in lines[i]):
                            i += 1
                else:
                    parts = re.split(r'\s{2,}|\t|：|:', line)
                    headers = [p.strip() for p in parts if p.strip() and len(p.strip()) > 0]
                    if not headers or len(headers) < 2:
                        headers = ['序号', '事项描述', '负责人', '备注/期望结果']
                    i += 1
                
                if not headers:
                    headers = ['序号', '事项描述', '负责人', '备注/期望结果']
                
                # 解析表格数据行
                current_row = []
                row_number = 1
                
                while i < len(lines):
                    data_line = lines[i].strip()
                    
                    if not data_line:
                        if current_row:
                            while len(current_row) < len(headers):
                                current_row.append('')
                            table_data.append(current_row[:len(headers)])
                            current_row = []
                        i += 1
                        if i < len(lines) and not lines[i].strip():
                            break
                        continue
                    
                    if '|' in data_line:
                        if data_line.startswith('|') and data_line.endswith('|'):
                            row_data = [d.strip() for d in data_line.split('|')[1:-1]]
                            if len(row_data) >= len(headers):
                                table_data.append(row_data[:len(headers)])
                            i += 1
                        else:
                            break
                    elif re.match(r'^\d+[\.、]', data_line):
                        if current_row:
                            while len(current_row) < len(headers):
                                current_row.append('')
                            table_data.append(current_row[:len(headers)])
                        
                        match = re.match(r'^(\d+)[\.、]\s*(.*)', data_line)
                        if match:
                            row_number = match.group(1)
                            rest = match.group(2).strip()
                            current_row = [row_number, rest]
                        else:
                            current_row = [str(row_number), data_line]
                        row_number += 1
                        i += 1
                    elif current_row and len(current_row) < len(headers):
                        is_new_column = False
                        if len(current_row) >= 2:
                            if any(keyword in data_line for keyword in ['发言人', '负责人', 'Speaker']):
                                is_new_column = True
                            elif len(current_row) >= 3 and any(keyword in data_line for keyword in ['评估', '备注', '期望', '结果']):
                                is_new_column = True
                        
                        if is_new_column or len(current_row) == 0:
                            current_row.append(data_line)
                        else:
                            last_col_idx = len(current_row) - 1
                            if last_col_idx >= 0:
                                current_row[last_col_idx] += '\n' + data_line
                        i += 1
                    else:
                        if current_row:
                            while len(current_row) < len(headers):
                                current_row.append('')
                            table_data.append(current_row[:len(headers)])
                        break
                
                if current_row:
                    while len(current_row) < len(headers):
                        current_row.append('')
                    table_data.append(current_row[:len(headers)])
                
                # 创建Word表格
                if table_data or headers:
                    try:
                        table = doc.add_table(rows=len(table_data) + 1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        set_table_borders_black(table)
                        
                        # 设置表头
                        header_row = table.rows[0]
                        for j, header in enumerate(headers):
                            if j < len(header_row.cells):
                                cell = header_row.cells[j]
                                cell.text = header
                                para = cell.paragraphs[0]
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in para.runs:
                                    run.bold = False
                                    run.font.name = '仿宋_GB2312'
                                    run.font.size = Pt(16)
                                    run.font.color.rgb = black_color
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                    run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                                    run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                        
                        # 填充数据
                        for row_idx, row_data in enumerate(table_data):
                            if row_idx + 1 < len(table.rows):
                                row = table.rows[row_idx + 1]
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < len(row.cells) and col_idx < len(headers):
                                        cell = row.cells[col_idx]
                                        cell.text = str(cell_data) if cell_data else ''
                                        para = cell.paragraphs[0]
                                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        for run in para.runs:
                                            run.font.name = '仿宋_GB2312'
                                            run.font.size = Pt(16)
                                            run.font.color.rgb = black_color
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                            run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                                            run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                    except Exception as e:
                        logger.error(f"创建表格失败: {e}")
                        para = doc.add_paragraph(line)
                        for run in para.runs:
                            run.font.name = '仿宋_GB2312'
                            run.font.size = Pt(16)
                            run.font.color.rgb = black_color
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                            run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                
                doc.add_paragraph()
                continue
            
            # 通用标题识别规则
            is_title = False
            title_text = line
            title_size = Pt(14)
            
            clean_line = re.sub(r'^\*\*|\*\*$', '', line).strip()
            
            # 规则1：Markdown格式的标题
            if re.match(r'^\*\*.*?\*\*', line):
                is_title = True
                title_text = clean_line
                if clean_line.endswith(':') or clean_line.endswith('：'):
                    title_size = Pt(12)
                else:
                    title_size = Pt(14)
            
            # 规则2：章节标题（中文数字开头）
            elif re.match(r'^[一二三四五六七八九十]+[、.]', clean_line):
                is_title = True
                title_text = clean_line
                title_size = Pt(16)
                current_level1_title = clean_line
                level2_counter = 0
            
            # 规则3：数字编号标题
            elif re.match(r'^\d+[、.]\s+', clean_line) and len(clean_line) > 3:
                is_title = True
                title_text = clean_line
                title_size = Pt(16)
            
            # 规则3.5：带括号的中文数字标题
            elif re.match(r'^[（(][一二三四五六七八九十]+[）)]', clean_line):
                is_title = True
                title_text = clean_line
                title_size = Pt(16)
            
            # 规则4：以冒号结尾的行
            elif clean_line.endswith(':') or clean_line.endswith('：'):
                if not re.search(r'[。，、]', clean_line) and len(clean_line) < 120:
                    title_part = clean_line.rstrip('：:').strip()
                    particle_count = len(re.findall(r'[的了在是就有]', title_part))
                    if particle_count < len(title_part) * 0.1:
                        is_title = True
                        if current_level1_title and not re.match(r'^[（(][一二三四五六七八九十]+[）)]', title_part):
                            level2_counter += 1
                            if level2_counter <= len(chinese_numerals):
                                title_text = f"（{chinese_numerals[level2_counter - 1]}）{clean_line}"
                            else:
                                title_text = clean_line
                        else:
                            title_text = clean_line
                        title_size = Pt(16)
            
            # 规则5：标题后跟内容的情况
            if not is_title:
                label_content_match = re.match(r'^(\*\*)?([^：:*]+?)(\*\*)?[:：]\s+(.+)$', line)
                if label_content_match:
                    potential_label = label_content_match.group(2).strip()
                    if (len(potential_label) < 100 and 
                        not re.search(r'[。，、]', potential_label)):
                        particle_count = len(re.findall(r'[的了在是就有]', potential_label))
                        if particle_count < len(potential_label) * 0.1:
                            is_title = True
                            clean_potential_label = re.sub(r'^\*\*|\*\*$', '', potential_label).strip()
                            content_part = label_content_match.group(4).strip()
                            
                            if current_level1_title and not re.match(r'^[（(][一二三四五六七八九十]+[）)]', clean_potential_label):
                                level2_counter += 1
                                if level2_counter <= len(chinese_numerals):
                                    numbered_label = f"（{chinese_numerals[level2_counter - 1]}）{clean_potential_label}"
                                    if line.startswith('**') and line.count('**') >= 2:
                                        title_text = f"**{numbered_label}**：{content_part}"
                                    else:
                                        title_text = f"{numbered_label}：{content_part}"
                                else:
                                    title_text = re.sub(r'^\*\*|\*\*$', '', line).strip()
                            else:
                                title_text = re.sub(r'^\*\*|\*\*$', '', line).strip()
                            title_size = Pt(16)
            
            if is_title:
                is_bracketed_title = re.match(r'^[（(][一二三四五六七八九十]+[）)]', title_text)
                is_chinese_num_title = re.match(r'^[一二三四五六七八九十]+[、.]', title_text)
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = 1.0
                is_small_title = (not is_chinese_num_title and not is_bracketed_title and 
                                 (title_text.endswith(':') or title_text.endswith('：')) and
                                 current_level1_title is not None)
                
                label_content_match = re.match(r'^([^：:]+?)[:：]\s+(.+)$', title_text)
                if label_content_match:
                    label_part = label_content_match.group(1).strip() + ': '
                    content_part = label_content_match.group(2).strip()
                    
                    label_run = para.add_run(label_part)
                    label_run.bold = False
                    if is_bracketed_title or is_small_title:
                        label_run.font.name = 'KaiTi_GB2312'
                        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                    elif is_chinese_num_title:
                        label_run.font.name = 'SimHei'
                        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    else:
                        label_run.font.name = 'SimHei'
                        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    label_run.font.size = title_size
                    label_run.font.color.rgb = black_color
                    
                    content_run = para.add_run(content_part)
                    content_run.bold = False
                    content_run.font.name = '仿宋_GB2312'
                    content_run.font.size = Pt(16)
                    content_run.font.color.rgb = black_color
                    content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    content_run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                    content_run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                else:
                    run = para.add_run(title_text)
                    run.bold = False
                    if is_bracketed_title or is_small_title:
                        run.font.name = 'KaiTi_GB2312'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                    elif is_chinese_num_title:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    else:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = title_size
                    run.font.color.rgb = black_color
            else:
                # 普通文本段落
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.15
                run = para.add_run(line)
                run.bold = False
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(16)
                run.font.color.rgb = black_color
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run._element.rPr.rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
            
            i += 1
        
        # ✅ 使用微秒级时间戳 + file_id 确保文件名唯一性
        now = datetime.now()
        timestamp = now.strftime('%Y%m%d_%H%M%S_%f')
        
        if file_id:
            file_id_short = file_id.replace('-', '')[:8]
            filename = f"{filename_prefix}_{timestamp}_{file_id_short}.docx"
        else:
            filename = f"{filename_prefix}_{timestamp}.docx"
        
        # 使用单独的会议纪要目录
        summary_dir = FILE_CONFIG.get('summary_dir', 'meeting_summaries')
        if not os.path.exists(summary_dir):
            os.makedirs(summary_dir, exist_ok=True)
            logger.info(f"创建会议纪要目录: {summary_dir}")
        
        filepath = os.path.join(summary_dir, filename)
        
        doc.save(filepath)
        return filename, filepath
        
    except Exception as e:
        logger.error(f"保存会议纪要Word文档失败: {e}")
        return None, None
